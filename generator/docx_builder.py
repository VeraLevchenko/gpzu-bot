# generator/docx_builder.py
from io import BytesIO
from typing import List, Optional
from copy import deepcopy
from pathlib import Path

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.table import Table, _Cell

from parsers.egrn_parser import EGRNData, Coord

# Базовые пути относительно проекта
BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = BASE_DIR / "templates" / "gpzu_template.docx"
TZ_DIR        = BASE_DIR / "templates" / "tz_reglament"

# Таблица координат (итого 17.88 см)
COL_W = [Cm(4.50), Cm(6.69), Cm(6.69)]
MARKER_COORDS = "[[COORDS_TABLE]]"
# Маркер для вставки файла по разделу 2.2
MARKER_TZ_INSERT = "[[INSERT_OD2_VRI]]"  # оставляем старое имя маркера, можно переименовать при желании

# ----------------- Утилиты таблицы координат ----------------- #
def _center_cell(cell: _Cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for par in cell.paragraphs:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _apply_table_layout(tbl: Table):
    try:
        tbl.autofit = False
    except Exception:
        pass
    try:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            if i < len(COL_W):
                try:
                    cell.width = COL_W[i]
                except Exception:
                    pass
            _center_cell(cell)

def _fmt_coord(v: Optional[str]) -> str:
    return (v or "").strip().replace(" ", "").replace(".", ",")

def _build_coords_table(doc: Document, coords: List[Coord]) -> Table:
    """
    Двухстрочная шапка:
      [0,0] "Обозначение..." (merge с [1,0])
      [0,1]-[0,2] общий заголовок коорд. (merge)
      [1,1] X, [1,2] Y
    """
    tbl = doc.add_table(rows=2, cols=3)
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass

    top = tbl.rows[0].cells
    bot = tbl.rows[1].cells

    top[0].text = "Обозначение (номер) характерной точки"
    top[1].text = ("Перечень координат характерных точек в системе координат, "
                   "используемой для ведения Единого государственного реестра недвижимости")
    top[2].text = ""
    bot[0].text = ""
    bot[1].text = "X"
    bot[2].text = "Y"

    top[0].merge(bot[0])   # вертикаль
    top[1].merge(top[2])   # горизонталь

    for c in (coords or []):
        r = tbl.add_row().cells
        r[0].text = (c.num or "").strip()
        r[1].text = _fmt_coord(c.x)
        r[2].text = _fmt_coord(c.y)

    _apply_table_layout(tbl)
    return tbl

# ----------------- Поиск и замены в документе ----------------- #
def _iter_all_paragraphs(doc: Document):
    """Итерировать параграфы в документе, включая те, что внутри таблиц."""
    for p in doc.paragraphs:
        yield p
    def walk_cell(cell: _Cell):
        for p in cell.paragraphs:
            yield p
        for t in cell.tables:
            for r in t.rows:
                for c in r.cells:
                    yield from walk_cell(c)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                yield from walk_cell(c)

def _find_paragraph_with_text(doc: Document, marker: str):
    for p in _iter_all_paragraphs(doc):
        if p.text and marker in p.text:
            return p
    return None

def _replace_paragraph_with_table(anchor_paragraph, table: Table):
    """Вставить таблицу сразу после параграфа и удалить сам параграф."""
    anchor_elm = anchor_paragraph._element
    parent = anchor_elm.getparent()
    parent.insert(parent.index(anchor_elm) + 1, table._tbl)
    parent.remove(anchor_elm)

def _insert_external_docx_at_paragraph(main_doc: Document, marker: str, ext_path: Path) -> bool:
    """
    Заменяет параграф с маркером содержимым другого DOCX (в то же место).
    Возвращает True, если вставка выполнена, иначе False (в т.ч. если файла нет).
    """
    p = _find_paragraph_with_text(main_doc, marker)
    if not p:
        return False
    parent = p._element.getparent()
    idx = parent.index(p._element)

    if not ext_path or not ext_path.exists():
        # Файла нет — просто удалим маркер, ничего не вставляя
        parent.remove(p._element)
        return False

    ext = Document(ext_path.as_posix())
    parent.remove(p._element)

    insert_pos = idx
    for el in list(ext.element.body):
        parent.insert(insert_pos, deepcopy(el))
        insert_pos += 1
    return True

def _remove_marker_paragraph(main_doc: Document, marker: str):
    """Удаляет параграф с данным маркером (если есть)."""
    p = _find_paragraph_with_text(main_doc, marker)
    if not p:
        return
    parent = p._element.getparent()
    parent.remove(p._element)

# ----------------- Основная функция генерации ----------------- #
def build_section1_docx(egrn: EGRNData, *, zone_name: str = "", zone_code: Optional[str] = None) -> bytes:
    """
    Рендер раздела 1:
      - фигурные плейсхолдеры docxtpl (в т.ч. {{ parcel.zone_name }});
      - подмена [[COORDS_TABLE]] реальной таблицей координат;
      - вставка внешнего DOCX по выбранному коду зоны:
          templates/tz_reglament/<КОД>_vri.docx
        если файла нет — маркер удаляется.
    """
    tpl = DocxTemplate(TEMPLATE_PATH.as_posix())

    ctx = {
        "gpzu": {"number": "", "application_ref": ""},
        "parcel": {
            "region": "Кемеровская область – Кузбасс",
            "municipality": "Новокузнецкий городской округ",
            "settlement": "",
            "cadnum": egrn.cadnum or "",
            "address": egrn.address or "",
            "area": egrn.area or "",
            "capital_objects_text": (
                ", ".join(egrn.capital_objects) if egrn.capital_objects
                else "Объекты капитального строительства отсутствуют"
            ),
            "zone_name": zone_name or "",   # {{ parcel.zone_name }}
        },
    }

    # 1) Рендерим шаблон docxtpl
    bio = BytesIO()
    tpl.render(ctx)
    tpl.save(bio)
    bio.seek(0)

    # 2) Постобработка через python-docx
    doc = Document(bio)

    # 2.1) Таблица координат вместо [[COORDS_TABLE]]
    p_coords = _find_paragraph_with_text(doc, MARKER_COORDS)
    if p_coords:
        tbl = _build_coords_table(doc, egrn.coordinates or [])
        _replace_paragraph_with_table(p_coords, tbl)

    # 2.2) Раздел 2.2 — подставляем файл по коду зоны, если он есть
    # Ищем файл строго по коду, например "Ж-1_vri.docx"; если не найден — пробуем в верхнем регистре
    inserted = False
    if zone_code:
        code_str = str(zone_code).strip()
        candidate_paths = [
            TZ_DIR / f"{code_str}_vri.docx",
            TZ_DIR / f"{code_str.upper()}_vri.docx",
        ]
        for path in candidate_paths:
            if path.exists():
                inserted = _insert_external_docx_at_paragraph(doc, MARKER_TZ_INSERT, path)
                break

    # Если ничего не вставили (файла нет или зона не выбрана) — удалим маркер
    if not inserted:
        _remove_marker_paragraph(doc, MARKER_TZ_INSERT)

    # 3) Сохраняем результат
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()
