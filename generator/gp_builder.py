import os
import logging
from copy import deepcopy
from pathlib import Path
from typing import Dict, Any, Optional, List

from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.table import Table, _Cell

logger = logging.getLogger("gpzu-bot.gp_builder")

# ----------------- Таблица координат (как в docx_builder.py) ----------------- #

COL_W = [Cm(4.50), Cm(6.69), Cm(6.69)]
MARKER_COORDS = "[[COORDS_TABLE]]"


def _center_cell(cell: _Cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for par in cell.paragraphs:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _apply_table_layout(tbl: Table):
    """Фиксированная ширина и выравнивание таблицы координат."""
    try:
        tbl.autofit = False
    except Exception:
        pass
    try:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            if i < len(COL_W):
                try:
                    cell.width = COL_W[i]
                except Exception:
                    pass
            _center_cell(cell)


def _fmt_coord(v: Optional[str]) -> str:
    """Формат числа: без пробелов, с запятой как разделителем."""
    return (v or "").strip().replace(" ", "").replace(".", ",")


def _iter_all_paragraphs(doc: Document):
    """Итерация по всем параграфам, включая те, что внутри таблиц."""
    for p in doc.paragraphs:
        yield p

    def walk_cell(cell: _Cell):
        for p in cell.paragraphs:
            yield p
        for t in cell.tables:
            for r in t.rows:
                for c in r.cells:
                    yield from walk_cell(c)

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                yield from walk_cell(c)


def _find_paragraph_with_text(doc: Document, marker: str):
    """Найти первый параграф, содержащий данный текст."""
    for p in _iter_all_paragraphs(doc):
        if p.text and marker in p.text:
            return p
    return None


def _replace_paragraph_with_table(anchor_paragraph, table: Table):
    """Вставить таблицу сразу после параграфа и удалить сам параграф."""
    anchor_elm = anchor_paragraph._element
    parent = anchor_elm.getparent()
    parent.insert(parent.index(anchor_elm) + 1, table._tbl)
    parent.remove(anchor_elm)


# --------------------------------------------------------------------------- #
#                             ОСНОВНОЙ КЛАСС                                   #
# --------------------------------------------------------------------------- #


class GPBuilder:
    """
    Генератор градостроительного плана земельного участка (ГПЗУ).

    Работает по шагам:
    1. Заполняет Jinja2-переменные в шаблоне `gpzu_template.docx`
    2. Вставляет блоки территориальных зон (ВРИ и параметры) из `data/tz_reglament`
    3. Заполняет таблицу ЗОУИТ и подставляет блоки ограничений из `data/zouit_reglament`
    4. Вставляет таблицу координат участка в место маркера `[[COORDS_TABLE]]`
    """

    def __init__(self, template_path: str, data_dir: Optional[str] = None):
        """
        Args:
            template_path: Путь к файлу шаблона gpzu_template.docx
            data_dir: Путь к папке `data` с подкаталогами `tz_reglament` и
                      `zouit_reglament`. Если None — определяется автоматически
                      от корня проекта.

        Структура проекта:

            gpzu-bot/
            ├── bot.py
            ├── generator/
            │   └── gp_builder.py   ← этот файл
            ├── templates/
            │   └── gpzu_template.docx
            └── data/
                ├── tz_reglament/
                └── zouit_reglament/
        """
        self.template_path = str(template_path)

        base_dir = Path(__file__).resolve().parent.parent  # .../gpzu-bot
        if data_dir is None:
            self.data_dir = base_dir / "data"
        else:
            self.data_dir = Path(data_dir)

        self.tz_dir = self.data_dir / "tz_reglament"
        self.zouit_dir = self.data_dir / "zouit_reglament"

        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Шаблон не найден: {self.template_path}")

        logger.info(f"GPBuilder: шаблон: {self.template_path}")
        logger.info(f"GPBuilder: data_dir: {self.data_dir}")
        logger.info(f"GPBuilder: tz_dir: {self.tz_dir}")
        logger.info(f"GPBuilder: zouit_dir: {self.zouit_dir}")

    # ------------------------------------------------------------------ #
    #   ЗОУИТ: подбор файла блока по названию (общая логика)
    # ------------------------------------------------------------------ #
    def get_zouit_block_filename(self, zouit_name: str) -> Optional[str]:
        """
        Определяет имя файла Word-блока ЗОУИТ по наименованию зоны.

        Файлы лежат в `data/zouit_reglament`:

            statia56_sanzona.docx        – санитарно-защитные зоны
            statia57_electro.docx        – охранные зоны объектов электросетевого хозяйства
            statia64_aeroport_full.docx  – приаэродромная территория (в целом)
            statia64_aeroport_4.docx     – четвертая подзона приаэродромной территории
        """
        name = (zouit_name or "").lower()

        # Санитарно-защитные зоны
        if "санитар" in name or "санитарно-защит" in name or "санитарно защит" in name:
            return "statia56_sanzona.docx"

        # Электросетевое хозяйство
        electro_keywords = [
            "охранная зона объектов электросетевого хозяйства",
            "охранная зона вл",
            "охранная зона кл",
            "электроэнергетики",
            "сооружение линейное электротехническое",
            "воздушной линии электропередачи",
            "электропередач",
        ]
        if any(k in name for k in electro_keywords):
            return "statia57_electro.docx"

        # Приаэродромная территория / подзоны
        if "приаэродром" in name or "аэродром" in name or "аэропорт" in name:
            if "четверт" in name and "подзон" in name:
                return "statia64_aeroport_4.docx"
            return "statia64_aeroport_full.docx"

        logger.warning(f"Не удалось подобрать файл ЗОУИТ по названию: {zouit_name!r}")
        return None

    # Совместимость с test_gp_builder.py
    def get_zouit_file(self, zouit_name: str) -> Optional[str]:
        """
        Обёртка для совместимости с test_gp_builder.py:
        возвращает имя файла блока ЗОУИТ (без пути) или None.
        """
        return self.get_zouit_block_filename(zouit_name)

    # ------------------------------------------------------------------ #
    #   Спец-логика выбора файла по ЗОУИТ (учёт номера и "четвертой")
    # ------------------------------------------------------------------ #
    def get_zouit_block_path(self, zouit: Dict[str, Any]) -> Optional[Path]:
        """
        Возвращает Path к файлу блока ЗОУИТ с учётом спец-логики:

        - если registry_number == "42:00-6.1695" → statia64_aeroport_full.docx
        - если в name есть слово "четвертая" → statia64_aeroport_4.docx
        - иначе используется стандартное сопоставление по названию
          (get_zouit_block_filename).
        """
        name = zouit.get("name") or ""
        registry_number = (zouit.get("registry_number") or "").strip()

        filename: Optional[str]

        # Специальный случай – приаэродромная территория целиком
        if registry_number == "42:00-6.1695":
            filename = "statia64_aeroport_full.docx"
        # Специальный случай – четвертая подзона
        elif "четверт" in name.lower():
            filename = "statia64_aeroport_4.docx"
        else:
            filename = self.get_zouit_block_filename(name)

        if not filename:
            return None

        path = self.zouit_dir / filename
        if not path.exists():
            logger.warning(f"Файл блока ЗОУИТ не найден: {path}")
            return None
        return path

    # ------------------------------------------------------------------ #
    #   Загрузка блоков зон (территориальные зоны)
    # ------------------------------------------------------------------ #
    def load_zone_block(self, zone_code: str, block_type: str) -> Optional[Document]:
        """
        Загружает Word-блок для территориальной зоны.

        Args:
            zone_code: Код зоны (например, 'Ж-1', 'ОД-1')
            block_type: 'vri'   – блок видов разрешенного использования (2.2)
                        'params' – блок параметров застройки (2.3)

        Ожидаемая структура файлов:

            data/tz_reglament/
                Ж-1_vri.docx   – ВРИ для Ж-1
                Ж-1.docx       – параметры для Ж-1
                ОД-1_vri.docx  – и т.д.
        """
        if block_type == "vri":
            filename = f"{zone_code}_vri.docx"
        else:
            filename = f"{zone_code}.docx"

        filepath = self.tz_dir / filename
        if not filepath.exists():
            logger.warning(f"Файл блока зоны не найден: {filepath}")
            return None

        logger.info(f"Загружен блок зоны: {filepath}")
        return Document(str(filepath))

    def load_zouit_block(self, zouit_name: str) -> Optional[Document]:
        """
        СТАРЫЙ интерфейс (по одному названию).
        Сейчас для раздела 5 лучше использовать get_zouit_block_path(zouit),
        т.к. там учитывается и реестровый номер.
        Оставлен для совместимости, может использоваться в тестах.
        """
        filename = self.get_zouit_block_filename(zouit_name)
        if not filename:
            return None

        filepath = self.zouit_dir / filename
        if not filepath.exists():
            logger.warning(f"Файл блока ЗОУИТ не найден: {filepath}")
            return None

        logger.info(f"Загружен блок ЗОУИТ (legacy): {filepath}")
        return Document(str(filepath))

    # ------------------------------------------------------------------ #
    #   Подготовка контекста для docxtpl
    # ------------------------------------------------------------------ #
    def prepare_context(self, gp_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Дополняет исходные данные служебными полями для шаблона.

        Добавляет:
        - capital_objects_text – текстовое описание объектов капстроительства
        - zouit_formatted      – список ЗОУИТ в удобном для шаблона формате
        - служебные переменные для сохранения маркеров вставки блоков
        """
        context = dict(gp_data)

        # Объекты капитального строительства
        capital_objects = gp_data.get("capital_objects") or []
        if capital_objects:
            parts: List[str] = []
            for idx, obj in enumerate(capital_objects, start=1):
                name = obj.get("name") or "Объект капитального строительства"
                area = obj.get("area")
                floors = obj.get("floors")

                obj_fragments: List[str] = [f"{idx}) {name}"]
                if area:
                    obj_fragments.append(f"площадью {area} кв. м")
                if floors:
                    obj_fragments.append(f"этажностью {floors} эт.")
                parts.append(", ".join(obj_fragments))

            context["capital_objects_text"] = "; ".join(parts)
        else:
            context["capital_objects_text"] = "Не предусмотрены"

        # ЗОУИТ в удобном виде для таблицы (раздел 6)
        zouit_raw = gp_data.get("zouit") or []
        formatted: List[Dict[str, str]] = []
        for z in zouit_raw:
            name = z.get("name") or ""
            registry_number = z.get("registry_number") or ""
            document = z.get("document") or ""
            restrictions = z.get("restrictions") or ""

            title = name
            if registry_number:
                title += f" ({registry_number})"

            formatted.append(
                {
                    "title": title,
                    "document": document,
                    "restrictions": restrictions,
                }
            )

        context["zouit_formatted"] = formatted

        # --- Сохранить маркеры вставки блоков после рендера Jinja ---
        # Иначе Jinja превратит {{INSERT_ZONE_VRI}} и другие в пустую строку,
        # и методы вставки блоков не найдут свои маркеры в документе.
        context["INSERT_ZONE_VRI"] = "{{INSERT_ZONE_VRI}}"
        context["INSERT_ZONE_PARAMS"] = "{{INSERT_ZONE_PARAMS}}"
        context["INSERT_ZOUIT_BLOCKS"] = "{{INSERT_ZOUIT_BLOCKS}}"

        return context

    # ------------------------------------------------------------------ #
    #   Вставка блоков по маркеру (территориальные зоны)
    # ------------------------------------------------------------------ #
    def insert_block_at_marker(self, doc: Document, marker: str, block_doc: Document) -> None:
        """
        Вставляет содержимое `block_doc` в документ `doc` на место абзаца,
        содержащего текст `marker` ({{INSERT_ZONE_VRI}} или {{INSERT_ZONE_PARAMS}}).
        """
        marker_para = None
        for para in doc.paragraphs:
            if marker in para.text:
                marker_para = para
                break

        if marker_para is None:
            logger.warning(f"Маркер {marker!r} не найден в документе")
            return

        # Удаляем сам маркер из текста абзаца
        marker_para.text = marker_para.text.replace(marker, "").strip()

        body = marker_para._p.getparent()
        idx = body.index(marker_para._p)

        # Клонируем элементы из блока (параграфы и таблицы) и вставляем
        elements = [deepcopy(el) for el in block_doc.element.body]
        for el in reversed(elements):
            body.insert(idx + 1, el)

    # ------------------------------------------------------------------ #
    #   Таблица ЗОУИТ (раздел 6)
    # ------------------------------------------------------------------ #
    def fill_zouit_table(self, doc: Document, zouit_list: List[Dict[str, Any]]) -> None:
        """
        Заполняет таблицу ЗОУИТ в разделе 6 градплана.

        Ищем таблицу, где в первой строке первой ячейки есть
        фраза «наименование зоны с особыми условиями».
        """
        if not zouit_list:
            logger.info("ЗОУИТ отсутствуют, таблица ЗОУИТ не заполняется")
            return

        target_table = None
        for table in doc.tables:
            if table.rows and table.rows[0].cells:
                first_cell_text = table.rows[0].cells[0].text.lower()
                if "наименование зоны с особыми условиями" in first_cell_text:
                    target_table = table
                    break

        if target_table is None:
            logger.warning("Таблица ЗОУИТ в документе не найдена")
            return

        # Удаляем все строки, кроме заголовка
        while len(target_table.rows) > 1:
            target_table._tbl.remove(target_table.rows[1]._tr)

        # Добавляем строки по каждой ЗОУИТ
        for z in zouit_list:
            name = z.get("name") or ""
            registry_number = z.get("registry_number") or ""
            area = z.get("area") or ""
            document = z.get("document") or ""
            restrictions = z.get("restrictions") or ""

            title = name
            if registry_number:
                title += f" ({registry_number})"

            row_cells = target_table.add_row().cells
            if len(row_cells) >= 1:
                row_cells[0].text = title
            if len(row_cells) >= 2:
                row_cells[1].text = area
            if len(row_cells) >= 3:
                row_cells[2].text = document
            if len(row_cells) >= 4:
                row_cells[3].text = restrictions

        logger.info(f"Таблица ЗОУИТ заполнена ({len(zouit_list)} записей)")

    # ------------------------------------------------------------------ #
    #   Вставка блоков ограничений ЗОУИТ (раздел 5)
    # ------------------------------------------------------------------ #
    def insert_zouit_blocks(self, doc: Document, zouit_list: List[Dict[str, Any]]) -> None:
        """
        Вставляет текстовые блоки ограничений для ЗОУИТ в раздел 5
        на место маркера {{INSERT_ZOUIT_BLOCKS}}.

        Для каждой ЗОУИТ формируется:

        - <НАИМЕНОВАНИЕ> (<РЕЕСТРОВЫЙ НОМЕР>), площадь земельного участка
          покрываемая зоной с особыми условиями использования территории
          составляет <area> кв.м;  — наименование и номер жирным,
        затем вставляется текст из соответствующего файла.

        Порядок — как в списке zouit_list.
        """
        marker = "{{INSERT_ZOUIT_BLOCKS}}"

        marker_para = None
        for para in doc.paragraphs:
            if marker in para.text:
                marker_para = para
                break

        if marker_para is None:
            logger.warning(f"Маркер {marker!r} не найден для вставки блоков ЗОУИТ")
            return

        # Убираем маркер из текста, но сам параграф оставляем как якорь
        marker_para.text = marker_para.text.replace(marker, "").strip()

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        body = marker_para._p.getparent()
        idx = body.index(marker_para._p)

        def add_header_paragraph(name: str, registry_number: str, area: str):
            """
            Добавляет абзац вида:
            - <ЖИРНОЕ НАЗВАНИЕ> (<ЖИРНЫЙ НОМЕР>), площадь земельного участка
              покрываемая зоной с особыми условиями использования территории
              составляет <area> кв.м;
            """
            nonlocal body, idx

            p = OxmlElement("w:p")

            # run 1: "- "
            r1 = OxmlElement("w:r")
            t1 = OxmlElement("w:t")
            t1.text = "- "
            r1.append(t1)
            p.append(r1)

            # run 2: жирное наименование (и номер, если есть)
            r2 = OxmlElement("w:r")
            rPr2 = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rPr2.append(b)
            r2.append(rPr2)

            t2 = OxmlElement("w:t")
            title = name
            if registry_number:
                title += f" ({registry_number})"
            t2.text = title
            r2.append(t2)
            p.append(r2)

            # run 3: хвост с площадью (обычным)
            if area:
                r3 = OxmlElement("w:r")
                t3 = OxmlElement("w:t")
                t3.text = (
                    ", площадь земельного участка покрываемая зоной с особыми "
                    "условиями использования территории составляет "
                    f"{area} кв.м;"
                )
                r3.append(t3)
                p.append(r3)

            body.insert(idx + 1, p)
            idx += 1

        for i, z in enumerate(zouit_list, start=1):
            name = z.get("name") or ""
            registry_number = (z.get("registry_number") or "").strip()
            area = (z.get("area") or "").strip()

            logger.info(f"Обработка ЗОУИТ {i}/{len(zouit_list)}: {name} ({registry_number})")

            # Заголовок с названием, номером и площадью
            add_header_paragraph(name, registry_number, area)

            # Определяем файл с текстом ограничений
            block_path = self.get_zouit_block_path(z)
            if block_path is None:
                # Если блока нет — вставляем предупреждение отдельным абзацем
                warn_p = OxmlElement("w:p")
                r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = (
                    f"[ВНИМАНИЕ: Не найден блок ограничений для ЗОУИТ "
                    f"'{name}' ({registry_number})]"
                )
                r.append(t)
                warn_p.append(r)
                body.insert(idx + 1, warn_p)
                idx += 1
                logger.warning(
                    f"Не найден файл блока ограничений для ЗОУИТ "
                    f"{name} ({registry_number})"
                )
            else:
                block_doc = Document(str(block_path))
                # Вставляем содержимое файла сразу после заголовка зоны
                elements = [deepcopy(el) for el in block_doc.element.body]
                for el in elements:
                    body.insert(idx + 1, el)
                    idx += 1

            # Пустой абзац между зонами
            if i < len(zouit_list):
                empty_p = OxmlElement("w:p")
                body.insert(idx + 1, empty_p)
                idx += 1

        logger.info(f"Вставлено блоков ЗОУИТ: {len(zouit_list)}")

    # ------------------------------------------------------------------ #
    #   Таблица координат участка
    # ------------------------------------------------------------------ #
    def insert_coords_table(self, doc: Document, coords: List[Dict[str, Any]]) -> None:
        """
        Вставляет таблицу координат земельного участка на место маркера
        [[COORDS_TABLE]].

        Ожидается список словарей:
            {"num": "1", "x": "2199812.21", "y": "438312.84"}

        Формат таблицы — как в docx_builder.py:
        - стиль "Table Grid" (границы),
        - двухстрочная шапка:
            1-я строка:
                [0,0] "Обозначение (номер) характерной точки"
                    (объединена с [1,0])
                [0,1]-[0,2] общий заголовок про перечень координат
            2-я строка:
                [1,1] "X", [1,2] "Y"
        - строки координат в том же порядке, как в coords.
        """
        if not coords:
            logger.info("Координаты отсутствуют, таблица координат не формируется")
            return

        # Ищем параграф с маркером [[COORDS_TABLE]]
        p_coords = _find_paragraph_with_text(doc, MARKER_COORDS)
        if not p_coords:
            logger.warning("Маркер [[COORDS_TABLE]] не найден, таблица координат не будет вставлена")
            return

        # Создаём таблицу: сначала 2 строки шапки
        tbl = doc.add_table(rows=2, cols=3)
        try:
            tbl.style = "Table Grid"  # границы таблицы
        except Exception:
            pass

        top = tbl.rows[0].cells
        bot = tbl.rows[1].cells

        # Первая строка шапки
        top[0].text = "Обозначение (номер) характерной точки"
        top[1].text = (
            "Перечень координат характерных точек в системе координат, "
            "используемой для ведения Единого государственного реестра недвижимости"
        )
        top[2].text = ""

        # Вторая строка шапки
        bot[0].text = ""
        bot[1].text = "X"
        bot[2].text = "Y"

        # Объединения ячеек как в образце:
        # [0,0] + [1,0] по вертикали
        top[0].merge(bot[0])
        # [0,1] + [0,2] по горизонтали
        top[1].merge(top[2])

        # Добавляем строки по координатам — строго в порядке из списка
        for coord in coords:
            r = tbl.add_row().cells
            r[0].text = str(coord.get("num") or "").strip()
            r[1].text = _fmt_coord(coord.get("x"))
            r[2].text = _fmt_coord(coord.get("y"))

        # Применяем ширины колонок и выравнивание
        _apply_table_layout(tbl)

        # Вставляем таблицу вместо параграфа с маркером
        _replace_paragraph_with_table(p_coords, tbl)

    # ------------------------------------------------------------------ #
    #   Основной метод генерации
    # ------------------------------------------------------------------ #
    def generate(self, gp_data: Dict[str, Any], output_path: str) -> str:
        """
        Генерирует градостроительный план.

        Шаги:
        1. Рендерим Jinja2-шаблон (DocxTemplate)
        2. Добавляем таблицу координат
        3. Вставляем блоки территориальной зоны (ВРИ, параметры)
        4. Заполняем таблицу ЗОУИТ и вставляем блоки ограничений
        5. Сохраняем итоговый файл
        """
        logger.info("Начало генерации градплана")

        # --- 1. Рендер шаблона через docxtpl ---
        tpl = DocxTemplate(self.template_path)
        context = self.prepare_context(gp_data)
        tpl.render(context)

        temp_path = str(Path(output_path).with_suffix(".tmp.docx"))
        tpl.save(temp_path)

        # Загружаем как обычный Document для низкоуровневых операций
        doc = Document(temp_path)

        # --- 2. Таблица координат ---
        parcel = gp_data.get("parcel") or {}
        coords = parcel.get("coordinates") or []
        if coords:
            self.insert_coords_table(doc, coords)
        else:
            logger.info("Координаты участка в данных отсутствуют")

        # --- 3. Блоки территориальных зон ---
        zone = gp_data.get("zone") or {}
        zone_code = zone.get("code")
        if zone_code:
            # ВРИ
            vri_block = self.load_zone_block(zone_code, "vri")
            if vri_block:
                self.insert_block_at_marker(doc, "{{INSERT_ZONE_VRI}}", vri_block)
            else:
                logger.warning(f"Не найден блок ВРИ для зоны {zone_code}")

            # Параметры
            params_block = self.load_zone_block(zone_code, "params")
            if params_block:
                self.insert_block_at_marker(doc, "{{INSERT_ZONE_PARAMS}}", params_block)
            else:
                logger.warning(f"Не найден блок параметров для зоны {zone_code}")

        # --- 4. ЗОУИТ ---
        zouit_list = gp_data.get("zouit") or []
        if zouit_list:
            self.fill_zouit_table(doc, zouit_list)
            self.insert_zouit_blocks(doc, zouit_list)
        else:
            logger.info("ЗОУИТ для участка отсутствуют")

        # --- 5. Сохранение результата ---
        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))

        # Удаляем временный файл
        try:
            os.remove(temp_path)
        except OSError:
            pass

        logger.info(f"ГПЗУ успешно сформирован: {out_path}")
        return str(out_path)


def generate_gp_document(gp_data: Dict[str, Any], output_path: str) -> str:
    """
    Утилитная функция для генерации ГПЗУ "в одно действие".

    Использует структуру проекта:

        gpzu-bot/
        ├── generator/gp_builder.py
        ├── templates/gpzu_template.docx
        └── data/...

    Поэтому путь к шаблону и data определяются относительно расположения
    текущего файла, а не текущей рабочей директории.
    """
    base_dir = Path(__file__).resolve().parent.parent
    template_path = base_dir / "templates" / "gpzu_template.docx"
    builder = GPBuilder(str(template_path))  # data_dir по умолчанию (base_dir / "data")
    return builder.generate(gp_data, output_path)
